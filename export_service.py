from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from datetime import datetime
from typing import Dict
import io

class DocumentExporter:
    def __init__(self):
        self.tum_blue = (0, 101, 189)  # TUM Corporate Blue

    def _create_filename(self, doc_type: str, extension: str) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_doc_type = doc_type.lower().replace(" ", "_")
        return f"TUM_{safe_doc_type}_{timestamp}.{extension}"

    def export_to_pdf(self, content: str, metadata: Dict[str, str]) -> bytes:
        """Export content to PDF and return bytes"""
        pdf = FPDF()
        pdf.add_page()
        
        # Header
        pdf.set_font("Arial", "B", 16)
        pdf.set_text_color(*self.tum_blue)
        pdf.cell(0, 10, f"TUM {metadata.get('doc_type', 'Document')}", ln=True, align="C")
        
        # Metadata
        pdf.set_font("Arial", "I", 10)
        pdf.set_text_color(128, 128, 128)
        pdf.cell(0, 10, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True)
        pdf.cell(0, 10, f"Tone: {metadata.get('tone', 'Standard')}", ln=True)
        pdf.ln(5)
        
        # Content
        pdf.set_font("Arial", "", 12)
        pdf.set_text_color(0, 0, 0)
        
        # Handle text encoding properly
        try:
            # Split content into lines and handle encoding
            lines = content.split('\n')
            for line in lines:
                # Encode to latin-1 for FPDF compatibility
                try:
                    encoded_line = line.encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 6, encoded_line)
                except:
                    # Fallback for problematic characters
                    safe_line = line.encode('ascii', 'replace').decode('ascii')
                    pdf.multi_cell(0, 6, safe_line)
                pdf.ln(2)
        except Exception as e:
            pdf.multi_cell(0, 6, f"Error displaying content: {str(e)}")
        
        # Return PDF as bytes
        return pdf.output(dest='S').encode('latin-1')

    def export_to_docx(self, content: str, metadata: Dict[str, str]) -> bytes:
        """Export content to DOCX and return bytes"""
        doc = Document()
        
        # Header
        header = doc.add_heading(f"TUM {metadata.get('doc_type', 'Document')}", level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Tone: {metadata.get('tone', 'Standard')}")
        doc.add_paragraph("=" * 50)
        
        # Content
        # Split content into paragraphs and add them properly
        paragraphs = content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_to_txt(self, content: str, metadata: Dict[str, str]) -> bytes:
        """Export content to TXT and return bytes"""
        text_content = f"""TUM {metadata.get('doc_type', 'Document')}
{'=' * 50}

Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}
Tone: {metadata.get('tone', 'Standard')}
{'=' * 50}

{content}
"""
        return text_content.encode('utf-8')

    def export_document(self, content: str, metadata: Dict[str, str], format: str) -> bytes:
        """Export document in specified format and return bytes"""
        if format == "pdf":
            return self.export_to_pdf(content, metadata)
        elif format == "docx":
            return self.export_to_docx(content, metadata)
        elif format == "txt":
            return self.export_to_txt(content, metadata)
        else:
            raise ValueError(f"Unsupported format: {format}")
